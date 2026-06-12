from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt

from backend.core.export.base import BaseExporter


class WordExporter(BaseExporter):
    def export(self, content: str, output_path: Path) -> Path:
        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Microsoft YaHei"
        font.size = Pt(11)

        for line in content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("######"):
                doc.add_heading(stripped.lstrip("# "), level=6)
            elif stripped.startswith("#####"):
                doc.add_heading(stripped.lstrip("# "), level=5)
            elif stripped.startswith("####"):
                doc.add_heading(stripped.lstrip("# "), level=4)
            elif stripped.startswith("###"):
                doc.add_heading(stripped.lstrip("# "), level=3)
            elif stripped.startswith("##"):
                doc.add_heading(stripped.lstrip("# "), level=2)
            elif stripped.startswith("#"):
                doc.add_heading(stripped.lstrip("# "), level=1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", stripped):
                text = re.sub(r"^\d+\.\s", "", stripped)
                doc.add_paragraph(text, style="List Number")
            elif stripped.startswith("> "):
                doc.add_paragraph(stripped[2:])
            elif stripped.startswith("```"):
                continue
            else:
                clean = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
                clean = re.sub(r"\*(.*?)\*", r"\1", clean)
                clean = re.sub(r"`(.*?)`", r"\1", clean)
                doc.add_paragraph(clean)

        doc.save(str(output_path))
        return output_path
